import streamlit as st
import google.generativeai as genai
from PIL import Image
import io
import os
import pandas as pd
from datetime import datetime

# 載入 Word 與 PDF、視覺化處理套件
try:
    from docx import Document
    from docx.shared import Inches, Cm
    from pypdf import PdfReader
    import plotly.express as px
except ImportError:
    st.error("❌ 雲端套件正在安裝中，請稍候一分鐘並重新整理網頁。")

# 1. 網頁風格設定 (日系療育風、馬卡龍色系)
st.set_page_config(page_title="幼兒學習發展分析系統", page_icon="🍃", layout="centered")

st.markdown("""
    <style>
    .stApp { background-color: #f0f7fa !important; }
    h1 { color: #5d707a !important; text-align: center; font-weight: bold; font-family: 'Comic Sans MS', cursive, sans-serif; } 
    div.stButton > button:first-child {
        background-color: #ffccbc !important; color: #5d4037 !important;
        border-radius: 25px !important; border: 2px solid #ffb6a0 !important;
        font-weight: bold !important; width: 100% !important; padding: 10px !important;
        box-shadow: 0px 4px 12px rgba(0,0,0,0.05) !important;
    }
    .stMarkdown { line-height: 1.8; color: #37474f; font-family: 'Verdana', sans-serif; }
    </style>
    """, unsafe_allow_html=True)

# 2. 安全讀取 Gemini API Key
if "GEMINI_API_KEY" in st.secrets:
    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
else:
    st.error("❌ 尚未在 Streamlit 後台設定 GEMINI_API_KEY")

try:
    available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
    target_model = next((m for m in available_models if 'gemini-1.5-flash' in m), available_models[0])
    model = genai.GenerativeModel(target_model)
except:
    model = genai.GenerativeModel('gemini-1.5-flash')

# 3. 智慧 PDF 課綱快取引擎
@st.cache_data
def load_curriculum_from_pdf(pdf_path):
    parsed_text = ""
    if os.path.exists(pdf_path):
        try:
            reader = PdfReader(pdf_path)
            for page in reader.pages:
                text = page.extract_text()
                if text: parsed_text += text + "\n"
            return parsed_text
        except Exception as e: return f"PDF 解析出錯: {e}"
    return "尚未上傳官方課綱 PDF 檔案"

pdf_file_path = os.path.join("data", "curriculum.pdf")
curriculum_raw_text = load_curriculum_from_pdf(pdf_file_path)

# 核心演算法：輸入生日，自動換算成「X歲Y個月」
def calculate_age_from_dob(dob_val):
    try:
        if pd.isna(dob_val) or not str(dob_val).strip(): return ""
        dob_str = str(dob_val).replace('/', '-').strip()
        for fmt in ("%Y-%m-%d", "%Y-%m-%d %H:%M:%S"):
            try:
                dob = datetime.strptime(dob_str.split()[0], fmt)
                break
            except: continue
        else:
            return "格式錯誤(請用YYYY-MM-DD)"
            
        today = datetime.now()
        years = today.year - dob.year
        months = today.month - dob.month
        if today.day < dob.day:
            months -= 1
        if months < 0:
            years -= 1
            months += 12
        return f"{years}歲{months}個月"
    except:
        return "計算失敗"

# 初始化 Session State 保險箱
if 'generated_report' not in st.session_state: st.session_state.generated_report = None
if 'current_child' not in st.session_state: st.session_state.current_child = ""
if 'current_age' not in st.session_state: st.session_state.current_age = ""
if 'roster' not in st.session_state: 
    st.session_state.roster = pd.DataFrame(columns=['座號', '幼兒姓名', '出生年月日 (例: 2022-05-20)'])

# ✨ 欄位完美升級：正式加入「完整指標鏈」儲存管線，支持同一次觀察有多重指標
if 'session_history' not in st.session_state: 
    st.session_state.session_history = pd.DataFrame(columns=['日期', '座號', '幼兒姓名', '幼兒年齡', '觀察類型', '對應領域', '主要指標', '完整指標鏈', '現有能力評估', '親師溝通'])

# 4. Word 檔案生成函式
def build_word_file(child_name, child_age, report_text, uploaded_images, template_file=None):
    report_data = {"活動紀錄": "", "領域分析": "", "能力評估": "", "智慧鷹架": "", "親師溝通": ""}
    sections = report_text.split('### ')
    for sec in sections:
        if '活動紀錄' in sec: report_data["活動紀錄"] = sec.split('\n', 1)[-1].strip().replace('**', '')
        elif '發展領域分析' in sec: report_data["領域分析"] = sec.split('\n', 1)[-1].strip().replace('**', '')
        elif '現有能力評估' in sec: report_data["能力評估"] = sec.split('\n', 1)[-1].strip().replace('**', '')
        elif '智慧鷹架引導' in sec: report_data["智慧鷹架"] = sec.split('\n', 1)[-1].strip().replace('**', '')
        elif '親師溝通筆記' in sec: report_data["親師溝通"] = sec.split('\n', 1)[-1].strip().replace('**', '')

    if template_file is not None:
        template_file.seek(0)
        doc = Document(template_file)
        replacements = {
            "{{幼兒姓名}}": child_name, "{{幼兒年齡}}": child_age,
            "{{活動紀錄}}": report_data["活動紀錄"], "{{領域分析}}": report_data["領域分析"],
            "{{能力評估}}": report_data["能力評估"], "{{親師溝通}}": report_data["親師溝通"], "{{智慧鷹架}}": report_data["智慧鷹架"]
        }
        for paragraph in doc.paragraphs:
            for tag, text in replacements.items():
                if tag in paragraph.text: paragraph.text = paragraph.text.replace(tag, text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for tag, text in replacements.items():
                        if tag in cell.text: cell.text = cell.text.replace(tag, text)
                    if "{{照片歷程}}" in cell.text:
                        cell.text = cell.text.replace("{{照片歷程}}", "")
                        img_paragraph = cell.paragraphs[0]
                        if uploaded_images:
                            for file in uploaded_images:
                                file.seek(0)
                                img_bytes = file.read()
                                img_stream = io.BytesIO(img_bytes)
                                try:
                                    img = Image.open(img_stream)
                                    img_width, img_height = img.size
                                    img_stream.seek(0)
                                    run = img_paragraph.add_run()
                                    if img_width > img_height: run.add_picture(img_stream, width=Cm(8), height=Cm(6))
                                    else: run.add_picture(img_stream, width=Cm(6), height=Cm(8))
                                    run.add_text("  ") 
                                except: pass
    else:
        doc = Document()
        doc.add_heading(f'🍃 {child_name} 專業報告', level=1)
        doc.add_paragraph(f'☀️ 幼兒年齡：{child_age}')
        doc.add_paragraph('------------------------------------------------------------------')
        lines = report_text.split('\n')
        for line in lines:
            cleaned = line.strip()
            if not cleaned: continue
            if line.startswith('### '): doc.add_heading(cleaned.replace('### ', '').replace('**', ''), level=2)
            elif line.startswith('#### '): doc.add_heading(cleaned.replace('#### ', '').replace('**', ''), level=3)
            elif line.startswith('##### '): doc.add_heading(cleaned.replace('##### ', '').replace('**', ''), level=4)
            elif line.startswith('* ') or line.startswith('- '): doc.add_paragraph(cleaned[2:].replace('**', ''), style='List Bullet')
            else: doc.add_paragraph(cleaned.replace('**', ''))
                
        if uploaded_images:
            doc.add_page_break() 
            doc.add_heading('📝 (照片歷程紀錄)', level=2)
            img_paragraph = doc.add_paragraph() 
            for idx, file in enumerate(uploaded_images):
                file.seek(0)
                img_bytes = file.read()
                img_stream = io.BytesIO(img_bytes)
                try:
                    img = Image.open(img_stream)
                    img_width, img_height = img.size
                    img_stream.seek(0) 
                    run = img_paragraph.add_run()
                    if img_width > img_height: run.add_picture(img_stream, width=Cm(8), height=Cm(6))
                    else: run.add_picture(img_stream, width=Cm(6), height=Cm(8))
                    run.add_text("  ")
                except: doc.add_paragraph(f"[圖片載入失敗]")
                    
    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream

# 5. 側邊欄：純檔案式大數據歷史載入區
with st.sidebar:
    st.markdown("### 📊 班級歷史紀錄檔案袋")
    st.write("數據皆存在您電腦本機中，系統不會留存任何隱私。")
    
    history_file = st.file_uploader("📂 請上傳您個人的歷史紀錄舊檔案 (.csv)", type=["csv"], key="hist_upload")
    
    if history_file is not None:
        try:
            uploaded_df = pd.read_csv(history_file)
            
            # 智慧老舊版本欄位自動相容補全
            if '觀察類型' not in uploaded_df.columns: uploaded_df['觀察類型'] = "學習區活動"
            if '主要指標' not in uploaded_df.columns and '學習指標' in uploaded_df.columns:
                uploaded_df['主要指標'] = uploaded_df['學習指標']
            if '完整指標鏈' not in uploaded_df.columns: uploaded_df['完整指標鏈'] = uploaded_df['主要指標']
            
            required_cols = ['日期', '座號', '幼兒姓名', '幼兒年齡', '觀察類型', '對應領域', '主要指標', '完整指標鏈', '現有能力評估', '親師溝通']
            # 濾除老舊的多餘欄位，維持乾淨結構
            uploaded_df = uploaded_df[required_cols]
            
            st.session_state.session_history = uploaded_df
            st.sidebar.success("🟢 班級歷史紀錄已智慧同步！")
        except Exception as e:
            st.sidebar.error(f"檔案讀取失敗: {e}")

    # 雙向度智慧大數據看板呈現
    if not st.session_state.session_history.empty:
        df_history = st.session_state.session_history
        
        # 面向一：六大領域平衡 (雷達圖 - 能力深度)
        st.markdown("---")
        st.markdown("### 📊 核心能力發展診斷 (六大領域)")
        domain_counts = df_history['對應領域'].value_counts().to_dict()
        all_domains = ["認知領域", "身體動作與健康領域", "語文領域", "社會領域", "情緒領域", "美感領域"]
        chart_data = {dom: domain_counts.get(dom, 0) for dom in all_domains}
        plot_df = pd.DataFrame(list(chart_data.items()), columns=['領域', '觀測次數'])
        fig_radar = px.line_polar(plot_df, r='觀測次數', theta='領域', line_close=True, color_discrete_sequence=['#ffccbc'])
        fig_radar.update_traces(fill='toself')
        fig_radar.update_layout(polar=dict(radialaxis=dict(visible=True)), showlegend=False, margin=dict(l=30, r=30, t=30, b=30), height=200)
        st.plotly_chart(fig_radar, use_container_width=True)
        
        # 面向二：教學情境投放診斷 (水平長條圖 - 環境作息)
        st.markdown("### 📂 教學情境分佈診斷 (觀察類型)")
        context_counts = df_history['觀察類型'].value_counts().to_dict()
        all_contexts = ["學習區活動", "自由遊戲", "戶外活動", "團體活動", "生活自理"]
        context_data = {ctx: context_counts.get(ctx, 0) for ctx in all_contexts}
        context_df = pd.DataFrame(list(context_data.items()), columns=['情境', '次數'])
        fig_bar = px.bar(context_df, x='次數', y='情境', orientation='h', color_discrete_sequence=['#b2dfdb'])
        fig_bar.update_layout(margin=dict(l=10, r=10, t=10, b=10), height=180, yaxis={'categoryorder':'total ascending'})
        st.plotly_chart(fig_bar, use_container_width=True)

# 6. 主畫面介面呈現 (大標題)
st.markdown("<h1>🍃 幼兒學習發展分析系統</h1>", unsafe_allow_html=True)
st.divider()

# 三連發分頁
tab_record, tab_chart, tab_roster = st.tabs(["📝 新增動態觀察", "📈 幼兒個人成長看板", "🎒 班級名單管理"])

# ---------------------------------------------------------
# 分頁三：🎒 班級名單管理
# ---------------------------------------------------------
with tab_roster:
    st.markdown("### 🎒 班級基本名單設定")
    uploaded_roster_file = st.file_uploader("📂 快速載入以前建好的班級名單 (.csv)", type=["csv"], key="roster_file_uploader")
    if uploaded_roster_file is not None:
        try:
            roster_df_import = pd.read_csv(uploaded_roster_file)
            required_roster_cols = ['座號', '幼兒姓名', '出生年月日 (例: 2022-05-20)']
            if all(col in roster_df_import.columns for col in required_roster_cols):
                st.session_state.roster = roster_df_import
                st.success("🟢 班級名單已智慧恢復！")
            else:
                st.error("❌ 上傳的名單格式不符。")
        except Exception as e:
            st.error(f"名單讀取失敗: {e}")

    st.write("---")
    st.write("📝 **線上編輯名單表格**：")
    edited_roster_df = st.data_editor(
        st.session_state.roster,
        num_rows="dynamic",
        use_container_width=True,
        key="roster_table_editor"
    )
    
    col_r1, col_r2 = st.columns(2)
    with col_r1:
        if st.button("💾 儲存並更新網頁選單"):
            st.session_state.roster = edited_roster_df
            st.success("🌸 網頁選單已即時連動更新！")
            
    with col_r2:
        csv_roster_buffer = io.StringIO()
        edited_roster_df.to_csv(csv_roster_buffer, index=False)
        roster_csv_bom = "\ufeff" + csv_roster_buffer.getvalue()
        st.download_button(label="📥 下載 / 備份本班級名單 (.csv)", data=roster_csv_bom, file_name="幼兒園班級名單_備份.csv", mime="text/csv")

# ---------------------------------------------------------
# 分頁一：📝 新增動態觀察
# ---------------------------------------------------------
with tab_record:
    if os.path.exists(pdf_file_path): st.caption("🟢 系統狀態：後台官方課綱 PDF 檔案已智慧載入。")
    else: st.caption("🟡 系統狀態：未偵測到後台 PDF，切換為 AI 常規跨領域推論模式。")

    roster_df = st.session_state.roster
    roster_names = roster_df['幼兒姓名'].dropna().tolist() if not roster_df.empty else []
    selectbox_options = ["請選擇班級幼兒... (或手動輸入)"] + roster_names + ["➕ 手動輸入新幼兒 (不在名單內)"]
    
    selected_name_option = st.selectbox("🖍️ 選擇幼兒姓名", options=selectbox_options)
    
    final_child_name = ""
    final_child_age = ""
    final_seat_num = "無"
    
    if selected_name_option == "請選擇班級幼兒... (或手動輸入)":
        st.caption("ℹ️ 請選擇上方的名單，或是下拉到最底下選擇手動輸入。")
    elif selected_name_option == "➕ 手動輸入新幼兒 (不在名單內)":
        col_manual_name, col_manual_age = st.columns(2)
        with col_manual_name: final_child_name = st.text_input("📝 請輸入全新幼兒姓名", placeholder="例如：強強")
        with col_manual_age: final_child_age = st.text_input("☀️ 請輸入年齡", placeholder="例如：4歲2個月")
    else:
        final_child_name = selected_name_option
        student_row = roster_df[roster_df['幼兒姓名'] == final_child_name].iloc[0]
        final_seat_num = str(student_row['座號'])
        dob_birthday = student_row['出生年月日 (例: 2022-05-20)']
        final_child_age = calculate_age_from_dob(dob_birthday)
        st.info(f"✨ 班級名單連動成功！【座號：{final_seat_num} 號】 ➜ 系統自動精算當前實齡：**{final_child_age}**")

    st.write("---")
    obs_type = st.selectbox(
        "🎯 選擇本次觀察類型 / 情境情境",
        ["學習區活動", "自由遊戲", "戶外活動", "團體活動", "生活自理"],
        help="選擇此紀錄發生的情境，系統將自動同步到左側的『環境投放診斷看板』中喔！"
    )

    teacher_notes = st.text_area("🖊️ 老師補充觀察 (為提升 AI 判定信心，建議補充以下資訊)", placeholder="1. 頻率：這是第一次發生、偶爾出現、還是穩定表現？\n2. 歷程：過程中是否失敗過？是否模仿同儕？\n3. 介入：老師有給予提示或動手幫忙嗎？\n4. 幼兒說的話：", key="input_notes")

    uploaded_files = st.file_uploader("🖼️ 1. 上傳觀察照片 (多張連拍歷程)", type=["jpg", "jpeg", "png"], accept_multiple_files=True)

    with st.expander("💡 第一次使用自訂模板？點我查看【專屬標籤設定教學】🌸"):
        st.markdown("""
        為了讓系統完美接軌您學校的行政格式，請在您電腦裡既有的 Word 空白紀錄表中，將需要讓系統代筆的地方，打上對應的「魔法標籤」：
        * 🖍️ **基本資訊**：`{{幼兒姓名}}`、`{{幼兒年齡}}`
        * 📝 **專業分析**：`{{活動紀錄}}`、`{{領域分析}}`、`{{能力評估}}`
        * 🌱 **支持策略**：`{{智慧鷹架}}`
        * 💌 **家長視角**：`{{親師溝通}}`
        * 🖼️ **動態相片牆**：`{{照片歷程}}`
        """)

    template_file_upload = st.file_uploader("📄 2. 上傳學校既有 Word 紀錄表模板 (選填)", type=["docx"])

    if st.button("🌟 生成專業觀察報告"):
        if uploaded_files and final_child_age and final_child_name:
            with st.spinner(f"正在透過證據權重模型與官方課綱 PDF 進行深度運算..."):
                try:
                    media_contents = []
                    # ✨ Prompt 升級：命令 AI 同時輸出「單一主指標（統計用）」與「全指標鏈（完整記錄）」
                    prompt = f"""
                    你是一套業界頂級的幼教發展評估可解釋性 AI (XAI) 系統。請基於「證據權重模型」綜合分析上傳的「系列照片 (動態歷程)」與以下資訊：
                    - 幼兒姓名：{final_child_name}
                    - 幼兒實齡：{final_child_age}
                    - 教師補充觀察：{teacher_notes if teacher_notes else "無"}

                    ⚠️ 【最高指導原則：教育部官方課綱絕對精確對照與核心欄位精準輸出】
                    系統後台已為你讀取了完整的教育部官方《幼兒園教保活動課程大綱》PDF 文本內容如下：
                    {curriculum_raw_text[:60000]} 

                    請依照以下結構輸出，語氣維持簡約療育風:

                    ### 🔍 【活動紀錄與行為證據】
                    客觀條列動態歷程證據。

                    ### 📊 【發展領域分析】
                    【核心領域標籤: (請從 PDF 課文中精確填入最具代表性的單一主領域，用於雷達圖：認知領域 / 身體動作與健康領域 / 語文領域 / 社會領域 / 情緒領域 / 美感領域)】
                    【核心指標標籤: (請精確寫出最核心的單一主指標編碼，例如：身-幼-2-1-1)】
                    【全指標鏈標籤: (請將本次行為所有符合的學習指標全部列出，並用直線符號隔開，例如：身-幼-2-1-1｜認-幼-2-1-2)】

                    - 符合核心課程目標：[核心主指標的完整字句]
                    - 本次統整性活動觸發之完整學習指標清單：
                      1. [指標一字句說明]
                      2. [指標二字句說明] (若有)
                      3. [指標三字句說明] (若有)

                    ### ✨ 【現有能力評估】
                    【判定結果】：(請擇一輸出：⭐ 正在建立 / ⭐⭐ 穩定運用 / ⭐⭐⭐ 靈活遷移)
                    【判定信心】：(請給出 0% - 100% 的估算值)
                    【主要依據】：簡述判定成立的核心觀察。
                    【尚缺證據】：指出若要提高信心分數還需要觀察到什麼。

                    第二段 (能力說明)：
                    描述對應發展常模的深層教育意義。

                    ### 🌱 【智慧鷹架引導】
                    ##### ☀️ ZPD 下一步推估
                    ##### ✨ 啟發性對話
                    ##### 📈 探索環境擴充

                    ### 🖍️ 【親師溝通筆記】
                    給家長看的內容 (約150字)，請執行「三步溝通法」保持手繪療育感。
                    """
                    media_contents.append(prompt)
                    for file in uploaded_files:
                        img = Image.open(file)
                        media_contents.append(img)
                    
                    response = model.generate_content(media_contents)
                    report_text = response.text
                    st.session_state.generated_report = report_text
                    st.session_state.current_child = final_child_name
                    st.session_state.current_age = final_child_age
                    
                    # 智慧解析與提取
                    extracted_domain = "未歸類"
                    extracted_indicator = "未識別"
                    extracted_chain = "未識別"
                    
                    if "【核心領域標籤:" in report_text: extracted_domain = report_text.split("【核心領域標籤:")[1].split("】")[0].strip()
                    if "【核心指標標籤:" in report_text: extracted_indicator = report_text.split("【核心指標標籤:")[1].split("】")[0].strip()
                    # ✨ 新增：提取完整的多重指標鏈字串
                    if "【全指標鏈標籤:" in report_text: 
                        extracted_chain = report_text.split("【全指標鏈標籤:")[1].split("】")[0].strip()
                    else:
                        extracted_chain = extracted_indicator
                    
                    eval_summary = "⭐ 正在建立"
                    if "⭐⭐慢 穩定運用" in report_text or "⭐⭐ 穩定運用" in report_text: eval_summary = "⭐⭐ 穩定運用"
                    elif "⭐⭐⭐ 靈活遷移" in report_text: eval_summary = "⭐⭐⭐ 靈活遷移"
                    
                    note_summary = report_text.split('### 🖍️ ')[-1].split('\n', 1)[-1].strip() if '### 🖍️ ' in report_text else "無"
                    
                    today_str = datetime.now().strftime("%Y-%m-%d")
                    new_row = pd.DataFrame([{
                        '日期': today_str, '座號': final_seat_num, '幼兒姓名': final_child_name, '幼兒年齡': final_child_age,
                        '觀察類型': obs_type, '對應領域': extracted_domain, '主要指標': extracted_indicator, '完整指標鏈': extracted_chain,
                        '現有能力評估': eval_summary, '親師溝通': note_summary[:100] + "..."
                    }])
                    st.session_state.session_history = pd.concat([st.session_state.session_history, new_row], ignore_index=True)
                    st.toast("🎉 本次觀察已智慧連動雙向度指標與完整指標鏈！", icon="📝")
                        
                except Exception as e:
                    st.error(f"系統發生錯誤，錯誤訊息：{e}")
        else:
            st.warning("請選取或輸入幼兒姓名、年齡並至少上傳一張照片喔！")

    if st.session_state.generated_report:
        st.subheader("📋 專業觀察分析報告")
        st.markdown(st.session_state.generated_report)
        st.divider()
        
        col_btn1, col_btn2 = st.columns(2)
        with col_btn1:
            word_data = build_word_file(st.session_state.current_child, st.session_state.current_age, st.session_state.generated_report, uploaded_files, template_file_upload)
            st.download_button(label=f"doc 匯出 {st.session_state.current_child} 的 Word 報告", data=word_data, file_name=f"{st.session_state.current_child}_觀察紀錄表.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
        with col_btn2:
            csv_buffer = io.StringIO()
            st.session_state.session_history.to_csv(csv_buffer, index=False)
            csv_data_with_bom = "\ufeff" + csv_buffer.getvalue()
            st.download_button(label="📈 下載包含座號的完整歷史紀錄 (.csv)", data=csv_data_with_bom, file_name=f"班級整合學習歷程_{datetime.now().strftime('%m%d')}.csv", mime="text/csv")

# ---------------------------------------------------------
# 分頁二：📈 幼兒個人成長看板
# ---------------------------------------------------------
with tab_chart:
    st.markdown("### 📈 幼兒長期發展成長圖表")
    if not st.session_state.session_history.empty:
        df_history = st.session_state.session_history
        unique_children = df_history['幼兒姓名'].unique().tolist()
        selected_child = st.selectbox("🖍️ 請選擇要調閱成長圖表的幼兒姓名：", unique_children)
        
        if selected_child:
            df_child = df_history[df_history['幼兒姓名'] == selected_child].copy()
            
            # 🔄 完美收合：將日期強制轉字串，根除 00:00:00.0032 微秒 Bug
            df_child['日期'] = df_child['日期'].astype(str)
            df_child = df_child.sort_values(by='日期')
            
            st.markdown(f"#### 🍃 {selected_child} 的縱向能力發展軌跡")
            
            status_map = {"⭐ 正在建立": 1, "⭐⭐ 穩定運用": 2, "⭐⭐⭐ 靈活遷移": 3}
            df_child['能力分數'] = df_child['現有能力評估'].map(status_map).fillna(1)
            
            # 🎨 核心繪圖引擎：依據「主要領域」多線分色呈現
            fig_line = px.line(
                df_child, x='日期', y='能力分數', color='對應領域', markers=True, text='現有能力評估',
                title=f"📈 {selected_child} 的跨領域能力演進圖 (各領域獨立分色)",
                color_discrete_sequence=['#ffccbc', '#b2dfdb', '#ffe082', '#c5cae9', '#e1bee7', '#ffcdd2']
            )
            fig_line.update_yaxes(tickvals=[1, 2, 3], ticktext=["正在建立", "穩定運用", "靈活遷移"], range=[0.5, 3.5])
            fig_line.update_traces(textposition="top center")
            fig_line.update_xaxes(type='category') # 鎖定 X 軸為分類格式，防微秒誤判
            st.plotly_chart(fig_line, use_container_width=True)
            
            # 展示包含「完整指標鏈」的質性大表格
            st.markdown("#### 📜 歷史觀察足跡摘要")
            st.dataframe(df_child[['日期', '座號', '幼兒姓名', '幼兒年齡', '觀察類型', '對應領域', '完整指標鏈', '現有能力評估', '親師溝通']], use_container_width=True, hide_index=True)
    else:
        st.info("💡 目前歷史資料庫為空。請先在左側上傳歷史紀錄，或開始新增動態觀察吧！")

st.markdown("<br><p style='text-align: center; color: #b0bec5;'>🍃 陪伴孩子在愛與探索中萌芽 🍃</p>", unsafe_allow_html=True)
